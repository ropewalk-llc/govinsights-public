#!/usr/bin/env python3
"""
Generate a styled Word document for GovInsights: The School Debt Timeline.
All content sourced from school-debt-timeline.html.

Usage: python3 generate_docx.py
Output: /Users/kim/Downloads/GovInsights_Chowan_School_Debt_Timeline.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# --- Color Constants ---
COLOR_BODY = RGBColor(0x1E, 0x29, 0x3B)
COLOR_H1 = RGBColor(0x0F, 0x17, 0x2A)
COLOR_H2 = RGBColor(0x1E, 0x29, 0x3B)
COLOR_BLUE = RGBColor(0x1E, 0x40, 0xAF)
COLOR_SUBTITLE = RGBColor(0x47, 0x55, 0x69)
COLOR_SOURCE = RGBColor(0x94, 0xA3, 0xB8)
COLOR_GRAY_MED = RGBColor(0x64, 0x74, 0x8B)
COLOR_DARK_BG = RGBColor(0x1E, 0x29, 0x3B)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_AMBER = RGBColor(0x92, 0x40, 0x0E)
COLOR_QUOTE = RGBColor(0x33, 0x41, 0x55)
COLOR_BLUE_DARK = RGBColor(0x1E, 0x3A, 0x5F)

OUTPUT_PATH = "/Users/kim/Downloads/GovInsights_Chowan_School_Debt_Timeline.docx"


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """Set cell borders. kwargs: top, bottom, left, right, each a dict with sz, val, color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
        tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="0" '
            f'w:color="{attrs.get("color", "auto")}"/>'
        )
        existing = tcBorders.find(qn(f'w:{edge}'))
        if existing is not None:
            tcBorders.remove(existing)
        tcBorders.append(element)


def set_paragraph_border_bottom(paragraph, color="DBEAFE", sz="12"):
    """Add bottom border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def set_paragraph_shading(paragraph, color_hex):
    """Set paragraph background shading."""
    pPr = paragraph._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{color_hex}"/>')
    pPr.append(shading)


def set_paragraph_border_left(paragraph, color="1E40AF", sz="18"):
    """Add left border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="4" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def set_paragraph_borders(paragraph, **kwargs):
    """Set multiple borders on a paragraph. kwargs: top, bottom, left, right."""
    pPr = paragraph._p.get_or_add_pPr()
    borders_xml = f'<w:pBdr {nsdecls("w")}>'
    for edge, attrs in kwargs.items():
        borders_xml += (
            f'<w:{edge} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", "4")}" w:space="{attrs.get("space", "4")}" '
            f'w:color="{attrs.get("color", "auto")}"/>'
        )
    borders_xml += '</w:pBdr>'
    pBdr = parse_xml(borders_xml)
    pPr.append(pBdr)


def add_run(paragraph, text, bold=False, italic=False, size=None, color=None, font_name="Calibri"):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    run.font.name = font_name
    return run


def add_body_paragraph(doc, text, bold=False, space_after=Pt(6)):
    """Add a standard body paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    add_run(p, text, bold=bold, size=11, color=COLOR_BODY)
    return p


def add_heading1(doc, text):
    """Add Heading 1 styled paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    add_run(p, text, bold=True, size=18, color=COLOR_H1)
    set_paragraph_border_bottom(p, color="DBEAFE", sz="12")
    return p


def add_heading2(doc, text):
    """Add Heading 2 styled paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, text, bold=True, size=13, color=COLOR_H2)
    return p


def add_source(doc, text):
    """Add source citation paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(0)
    add_run(p, text, size=9, color=COLOR_SOURCE)
    return p


def add_blockquote(doc, text, source_text=None):
    """Add a blockquote with left blue border."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    set_paragraph_border_left(p, color="1E40AF", sz="18")
    set_paragraph_shading(p, "F8FAFC")
    add_run(p, text, italic=True, size=11, color=COLOR_QUOTE)
    if source_text:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.5)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        set_paragraph_border_left(p2, color="1E40AF", sz="18")
        set_paragraph_shading(p2, "F8FAFC")
        add_run(p2, source_text, size=9, color=COLOR_SOURCE)
    return p


def add_note(doc, title, lines):
    """Add an editorial/context note with amber left border."""
    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(4) if i > 0 else Pt(8)
        p.paragraph_format.space_after = Pt(2) if i < len(lines) - 1 else Pt(8)
        set_paragraph_shading(p, "FFFBEB")
        set_paragraph_borders(p,
            left={"color": "F59E0B", "sz": "18", "space": "6"},
            top={"color": "FDE68A", "sz": "4" if i == 0 else "0", "val": "single" if i == 0 else "none"},
            bottom={"color": "FDE68A", "sz": "4" if i == len(lines) - 1 else "0", "val": "single" if i == len(lines) - 1 else "none"},
            right={"color": "FDE68A", "sz": "4", "val": "single"},
        )
        if i == 0:
            add_run(p, title, bold=True, size=9, color=COLOR_AMBER)
            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.15)
            p2.paragraph_format.space_before = Pt(2)
            p2.paragraph_format.space_after = Pt(2)
            set_paragraph_shading(p2, "FFFBEB")
            set_paragraph_borders(p2,
                left={"color": "F59E0B", "sz": "18", "space": "6"},
                right={"color": "FDE68A", "sz": "4", "val": "single"},
            )
            add_run(p2, line, size=9.5, color=COLOR_BODY)
        else:
            add_run(p, line, size=9.5, color=COLOR_BODY)


def add_summary_box(doc, paragraphs, premise=None):
    """Add summary box with blue background."""
    for i, text in enumerate(paragraphs):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(6) if i > 0 else Pt(8)
        p.paragraph_format.space_after = Pt(6) if i < len(paragraphs) - 1 else Pt(4)
        set_paragraph_shading(p, "EFF6FF")
        set_paragraph_borders(p,
            left={"color": "BFDBFE", "sz": "4"},
            right={"color": "BFDBFE", "sz": "4"},
            top={"color": "BFDBFE", "sz": "4" if i == 0 else "0", "val": "single" if i == 0 else "none"},
            bottom={"color": "BFDBFE", "sz": "4" if i == len(paragraphs) - 1 and not premise else "0",
                     "val": "single" if i == len(paragraphs) - 1 and not premise else "none"},
        )
        add_run(p, text, size=11, color=COLOR_BLUE_DARK)
    if premise:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        set_paragraph_shading(p, "EFF6FF")
        set_paragraph_borders(p,
            left={"color": "BFDBFE", "sz": "4"},
            right={"color": "BFDBFE", "sz": "4"},
            bottom={"color": "BFDBFE", "sz": "4"},
        )
        add_run(p, premise, bold=True, size=10, color=COLOR_BLUE)


def add_findings_box(doc, paragraphs):
    """Add findings box with gray background."""
    for i, text in enumerate(paragraphs):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.right_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(6) if i > 0 else Pt(8)
        p.paragraph_format.space_after = Pt(6) if i < len(paragraphs) - 1 else Pt(8)
        set_paragraph_shading(p, "F8FAFC")
        set_paragraph_borders(p,
            left={"color": "E2E8F0", "sz": "4"},
            right={"color": "E2E8F0", "sz": "4"},
            top={"color": "E2E8F0", "sz": "4" if i == 0 else "0", "val": "single" if i == 0 else "none"},
            bottom={"color": "E2E8F0", "sz": "4" if i == len(paragraphs) - 1 else "0",
                     "val": "single" if i == len(paragraphs) - 1 else "none"},
        )
        add_run(p, text, size=11, color=COLOR_BODY)


def add_bullet(doc, text, bold_prefix=None, size=11, indent=Inches(0.5)):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = indent
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    if bold_prefix:
        add_run(p, bold_prefix, bold=True, size=size, color=COLOR_BODY)
        add_run(p, text, size=size, color=COLOR_BODY)
    else:
        add_run(p, text, size=size, color=COLOR_BODY)
    return p


def create_table(doc, headers, rows, col_widths=None, highlight_rows=None, likely_col=None):
    """Create a styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        add_run(p, header, bold=True, size=8.5, color=COLOR_WHITE)
        set_cell_shading(cell, "1E293B")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        is_highlight = highlight_rows and r_idx in highlight_rows
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)

            # Check for LIKELY badge
            if "[LIKELY]" in str(cell_text):
                parts = str(cell_text).split("[LIKELY]")
                add_run(p, parts[0], bold=is_highlight, size=9.5, color=COLOR_BODY)
                add_run(p, " [LIKELY]", bold=True, size=9, color=COLOR_BLUE)
            else:
                bold = is_highlight
                add_run(p, str(cell_text), bold=bold, size=9.5, color=COLOR_BODY)

            # Right-align numeric columns
            if cell_text and (str(cell_text).startswith("$") or str(cell_text).startswith("~$")
                             or str(cell_text).startswith("+$") or str(cell_text).startswith("+")
                             or str(cell_text).startswith("n/a")):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Row shading
            if is_highlight:
                set_cell_shading(cell, "EFF6FF")
                if c_idx == 0:
                    set_cell_border(cell, left={"color": "1E40AF", "sz": "18"})
            elif r_idx % 2 == 1:
                set_cell_shading(cell, "F8FAFC")

    # Set column widths
    if col_widths:
        for r_idx, row in enumerate(table.rows):
            for c_idx, width in enumerate(col_widths):
                row.cells[c_idx].width = width

    # Set cell padding
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                f'  <w:top w:w="40" w:type="dxa"/>'
                f'  <w:left w:w="80" w:type="dxa"/>'
                f'  <w:bottom w:w="40" w:type="dxa"/>'
                f'  <w:right w:w="80" w:type="dxa"/>'
                f'</w:tcMar>'
            )
            tcPr.append(tcMar)

    return table


def add_foia_item(doc, description, target):
    """Add a FOIA item with amber left border."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    set_paragraph_border_left(p, color="F59E0B", sz="12")
    add_run(p, description, size=11, color=COLOR_BODY)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.2)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)
    set_paragraph_border_left(p2, color="F59E0B", sz="12")
    add_run(p2, target, size=8.5, color=COLOR_AMBER)


def add_page_break(doc):
    """Add a page break."""
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def setup_header_footer(doc, section, is_cover=False):
    """Set up headers and footers for a section."""
    if is_cover:
        section.different_first_page_header_footer = True
        # First page (cover) has no header/footer
        header = section.first_page_header
        header.is_linked_to_previous = False
        for p in header.paragraphs:
            p.text = ""
        footer = section.first_page_footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.text = ""

    # Regular header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.clear()

    # Use a tab to separate left and right content
    tab_stops = hp.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.8), WD_ALIGN_PARAGRAPH.RIGHT)

    add_run(hp, "GovInsights", bold=True, size=9, color=COLOR_BLUE)
    add_run(hp, "\t", size=9)
    add_run(hp, "The School Debt Timeline, Chowan County, NC", size=9, color=COLOR_SOURCE)

    # Bottom border on header
    set_paragraph_border_bottom(hp, color="E2E8F0", sz="8")

    # Footer with page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page number field
    run = fp.add_run()
    run.font.size = Pt(9)
    run.font.color.rgb = COLOR_SOURCE
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = fp.add_run()
    run2.font.size = Pt(9)
    run2.font.color.rgb = COLOR_SOURCE
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run3 = fp.add_run()
    run3.font.size = Pt(9)
    run3.font.color.rgb = COLOR_SOURCE
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)


def build_document():
    doc = Document()

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.9)
    section.different_first_page_header_footer = True

    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = COLOR_BODY

    # Fix List Bullet style
    if 'List Bullet' in doc.styles:
        lb = doc.styles['List Bullet']
        lb.font.name = 'Calibri'
        lb.font.size = Pt(11)
        lb.font.color.rgb = COLOR_BODY

    # Cover page: blank header/footer
    first_header = section.first_page_header
    first_header.is_linked_to_previous = False
    if first_header.paragraphs:
        first_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    if first_footer.paragraphs:
        first_footer.paragraphs[0].text = ""

    # Set up regular header/footer
    setup_header_footer(doc, section, is_cover=True)

    # ================================================================
    # COVER PAGE
    # ================================================================

    # Spacer
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    # GovInsights logo text
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    add_run(p, "GovInsights", bold=True, size=24, color=COLOR_BLUE)

    # Blue line
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run()
    # Create a small colored line using paragraph border
    set_paragraph_border_bottom(p, color="1E40AF", sz="24")

    # Title
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    add_run(p, "The School Debt Timeline", bold=True, size=28, color=COLOR_H1)

    # Subtitle
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    add_run(p, "How Chowan County built an $85 Million high school after voters voted No on an earlier referendum", size=14, color=COLOR_SUBTITLE)

    # Description
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "How decisions about a new high school connect to your property taxes.", size=11, color=COLOR_SUBTITLE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "An analysis of public minutes and documents from the", size=11, color=COLOR_SUBTITLE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(30)
    add_run(p, "Chowan County Board of Commissioners.", size=11, color=COLOR_SUBTITLE)

    # Metadata
    meta_items = [
        ("Location: ", "Chowan County, North Carolina"),
        ("Period Analyzed: ", "November 2018 to March 2026"),
        ("Documents Reviewed: ", "480 BOC minutes and agenda packets"),
        ("Prepared by: ", "Ropewalk Technologies LLC"),
        ("Published: ", "April 2026"),
    ]
    for label, value in meta_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        add_run(p, label, bold=True, size=9, color=COLOR_GRAY_MED)
        add_run(p, value, size=9, color=COLOR_SOURCE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "All quotes are verbatim from official Chowan County Board of Commissioner meeting minutes.", size=9, color=COLOR_SOURCE)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "All minutes are publicly available at chowancounty-nc.gov.", size=9, color=COLOR_SOURCE)

    # Page break after cover
    add_page_break(doc)

    # ================================================================
    # SUMMARY
    # ================================================================
    add_heading1(doc, "Summary")

    add_summary_box(doc, [
        'Among the most expensive high schools per student ever built in rural America, John A. Holmes High School cost $153,430 per student for 554 students. This is 2.5 to 7.5 times the national typical range of $20,000 to $60,000 per student for suburban school construction.',
        'When your property tax bill arrives, it can be hard to understand why it changed. The decisions that determine your taxes are spread across years of board meetings, budget votes, and financing agreements. Each one may seem routine on its own, but together they form a chain that directly affects what you pay.',
        'This report traces that chain for Chowan County. It follows the sequence of decisions surrounding this new high school, from the 2018 bond referendum through the 2026 property reappraisal, using the county\'s own meeting minutes and budget documents. These decisions may appear unconnected, but they are directly linked when it comes to how much debt the county carries and how that debt is paid for through your property taxes.',
        'The goal is clarity. County budgets are complex, and most residents do not have time to read seven years of board minutes. This analysis does that work and presents the documented sequence so that citizens can ask informed questions at budget hearings and understand what is driving changes to their tax bills.',
    ], premise='Informed citizens make up the backbone of local governance. Access to factual information about how public resources are used produces measurably better governance outcomes. That is the entire premise of GovInsights.')

    # ================================================================
    # DETAILED FINDINGS
    # ================================================================
    add_heading2(doc, "Detailed Findings")

    add_findings_box(doc, [
        'In November 2018, Chowan County voters rejected a bond referendum for a new John A. Holmes High School. Under the North Carolina Constitution, general obligation bonds require voter approval, and the voters said no.',

        'A second referendum was planned for 2020 but cancelled due to COVID-19. The Board subsequently approved USDA installment financing under G.S. 160A-20, which does not require voter approval. The guidance that no referendum was needed came not from the county attorney or the Local Government Commission, but from Rick Ott of M.B. Kahn Construction, a firm that was already working on the project and receiving county payments at the time.',

        'County expenditure records show that M.B. Kahn received $404,171 in eight payments from county Fund 45 between January and August 2021, before the Board of Commissioners voted on the Kahn contract on September 20, 2021. Who authorized these pre-vote payments is not documented in the 480 BOC documents reviewed for this report. Under NC G.S. \u00a7159-28, county expenditures require prior authorization. No hiring resolution, contract award, or qualifications-based selection process (as required for CMAR under G.S. 143-128.1) was found in any BOC document. Total payments to Kahn through August 2023 reached $1,407,511.',

        'The project, originally scoped at $50 million with a Davenport & Associates projection of "no impact on current tax rate," grew to approximately $85 million through a series of Board-approved scope and budget changes. M.B. Kahn\'s Construction Manager at Risk (CMAR) fee was set at 3.5% of construction cost, originally scoped on a $40 million base ($1.4 million). As costs grew, the fee grew with it. Chair Kirby confronted Kahn on record about the fee increase; Kahn\'s representative responded that the original commitment referred to the percentage, not the dollar amount.',

        'Separately from the school project, the county has never adopted a revenue-neutral tax rate after a reappraisal. In 2022, after the first reappraisal since the school debt was approved, the Board adopted a rate 4.7% above revenue-neutral and raised the rate again in 2024. Also in 2022, the Board shortened the reappraisal cycle from 8 years to 4 years. The 2026 reappraisal, which took effect January 1, 2026, increased property values county-wide by an estimated 38-42%, with individual increases ranging from 15% to 161%.',

        'Based on the Board\'s documented pattern of setting rates above revenue-neutral, homeowners can expect a cumulative tax increase of approximately 18-30% since FY2021-22.',
    ])

    add_source(doc, "All claims sourced to official Chowan County BOC meeting minutes and agenda packets, with supporting data from the U.S. Census Bureau, NC DPI, NC DOR, and the UNC School of Government.")

    # ================================================================
    # THE DOCUMENTED RECORD
    # ================================================================
    add_page_break(doc)
    add_heading1(doc, "The Documented Record")

    # 2018
    add_heading2(doc, "2018: Voters Reject the School Bond")
    add_body_paragraph(doc, "November 2018. Chowan County voters reject a bond referendum for a new John A. Holmes High School. Under the NC Constitution (Art. V, Sec. 4), general obligation bonds require voter approval. The voters said no.")

    # 2019
    add_heading2(doc, "2019: The County Prepares to Try Again")
    add_body_paragraph(doc, "November 18, 2019. The Board approves a financial advisor (Davenport & Company, up to $30,000) and bond counsel (Parker Poe, ~$20,000) to prepare for a 2020 bond referendum.")
    add_blockquote(doc,
        '"Commissioner McLaughlin objects that the process \'seems rushed.\'"',
        "Source: BOC November 18, 2019 Minutes")

    # 2020
    add_heading2(doc, "2020: COVID Cancels the Second Referendum")
    add_body_paragraph(doc, "May 4-18, 2020. The Board abandons the 2020 bond referendum due to COVID-19.")
    add_blockquote(doc,
        '"Due to the impact of COVID 19, we would no longer pursue a bond referendum for 2020."',
        "Chair Kersey, BOC May 2020 Minutes")

    # Featured: August 26, 2020
    # Gray background section
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    set_paragraph_shading(p, "F1F5F9")
    set_paragraph_borders(p,
        top={"color": "CBD5E1", "sz": "4"},
        left={"color": "CBD5E1", "sz": "4"},
        right={"color": "CBD5E1", "sz": "4"},
    )
    add_run(p, "August 26, 2020: The Meeting That Changed the Path", bold=True, size=13, color=COLOR_H1)

    featured_paras = [
        "Special meeting. Mr. Ott (M.B. Kahn), architect Paul Boney, County Manager Howard, and Superintendent Dr. Sasscer present school construction plans. All seven commissioners present.",
        "At this date, a November 2022 bond referendum was still the stated plan:",
    ]
    for text in featured_paras:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.15)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        set_paragraph_shading(p, "F1F5F9")
        set_paragraph_borders(p,
            left={"color": "CBD5E1", "sz": "4"},
            right={"color": "CBD5E1", "sz": "4"},
        )
        add_run(p, text, size=11, color=COLOR_BODY)

    add_blockquote(doc,
        '"He stated that the bond referendum is scheduled for November of 2022."',
        "Dr. Sasscer, August_26_Minutes.pdf, lines 196-197")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    set_paragraph_shading(p, "F1F5F9")
    set_paragraph_borders(p,
        left={"color": "CBD5E1", "sz": "4"},
        right={"color": "CBD5E1", "sz": "4"},
    )
    add_run(p, "Mr. Ott highlights the risk of the bond path, using Camden County as an example:", size=11, color=COLOR_BODY)

    add_blockquote(doc,
        '"He noted that if their bond referendum does not pass, Camden County will have to return those monies to the State."',
        "Mr. Ott, August_26_Minutes.pdf, line 88")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, "Context: Ott was illustrating that the bond referendum path carried risk of voter rejection (as Chowan had already experienced in 2018). This framed USDA installment financing, which requires no voter approval, as the safer alternative.", size=9, color=COLOR_SOURCE)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    set_paragraph_shading(p, "F1F5F9")
    set_paragraph_borders(p,
        left={"color": "CBD5E1", "sz": "4"},
        right={"color": "CBD5E1", "sz": "4"},
    )
    add_run(p, "Then, in the final four lines before adjournment:", size=11, color=COLOR_BODY)

    add_blockquote(doc,
        '"Stated the Commissioners intention is to move forward but the time frame needs to be slid a little bit. He asked if a referendum is needed to apply for USDA loan."',
        "Commissioner Kirby, August_26_Minutes.pdf, lines 201-202")
    add_blockquote(doc,
        '"Stated no that is not required."',
        "Mr. Ott (Construction Manager, M.B. Kahn), August_26_Minutes.pdf, line 203")
    add_blockquote(doc,
        '"Noted all of this has to be approved by the LGC."',
        "Commissioner Kirby, August_26_Minutes.pdf, line 204")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_after = Pt(8)
    set_paragraph_shading(p, "F1F5F9")
    set_paragraph_borders(p,
        left={"color": "CBD5E1", "sz": "4"},
        right={"color": "CBD5E1", "sz": "4"},
        bottom={"color": "CBD5E1", "sz": "4"},
    )
    add_run(p, "Meeting adjourned.", italic=True, size=11, color=COLOR_BODY)

    # Context note
    add_note(doc, "Context", [
        "No commissioner asked a follow-up question. No one asked the county attorney to verify.",
        "The guidance that no referendum was needed came from the project's own construction contractor, not from independent legal counsel, not from the county attorney, and not from the Local Government Commission.",
        "Under NC law (G.S. \u00a7 160A-20), installment financing such as USDA loans does not require voter approval. The legal statement was accurate. However, the source of the guidance was a party that stood to benefit from the project proceeding.",
        "M.B. Kahn had been introduced as \"the school's construction management firm\" at the January 14, 2020 Joint Committee meeting, 20 months before the BOC voted on their contract.",
    ])

    # ================================================================
    # 2021 (page break)
    # ================================================================
    add_page_break(doc)
    add_heading2(doc, "2021: USDA Financing Approved, Kahn Payments Begin")

    add_body_paragraph(doc, "Kahn was hired by the Board of Education for preconstruction and planning services prior to the BOC vote. County Fund 45 payments to Kahn began January 2021, eight months before the BOC voted on the contract in September 2021.")
    add_body_paragraph(doc, 'The authorization for those county fund payments prior to the BOC vote is not documented in BOC records. Was there a separate BOE authorization that covered the preconstruction work? Was Fund 45 set up with BOE money transferred to the county? The records reviewed for this report do not answer this.')

    # Kahn pre-contract payments table
    kahn_headers = ["Date", "Payee", "Amount", "Note"]
    kahn_note = "All 8 payments made before the Sep 20, 2021 BOC vote on the Kahn contract. Authorization source not documented in BOC records."
    kahn_rows = [
        ["1/14/2021", "MB Kahn", "$59,063", kahn_note],
        ["1/14/2021", "MB Kahn", "$43,019", ""],
        ["2/4/2021", "MB Kahn", "$40,298", ""],
        ["3/11/2021", "MB Kahn", "$39,696", ""],
        ["4/8/2021", "MB Kahn", "$44,153", ""],
        ["5/13/2021", "MB Kahn", "$44,293", ""],
        ["6/30/2021", "MB Kahn", "$39,450", ""],
        ["8/12/2021", "MB Kahn", "$94,199", ""],
    ]

    table = doc.add_table(rows=10, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(kahn_headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        add_run(p, h, bold=True, size=8.5, color=COLOR_WHITE)
        set_cell_shading(cell, "1E293B")
        if i == 2:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Data rows 1-8
    for r_idx in range(8):
        row_data = kahn_rows[r_idx]
        for c_idx in range(3):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_run(p, row_data[c_idx], size=9.5, color=COLOR_BODY)
            if c_idx == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F8FAFC")

        # Note column: merge rows 1-8
        note_cell = table.rows[r_idx + 1].cells[3]
        if r_idx % 2 == 1:
            set_cell_shading(note_cell, "F8FAFC")

    # Merge note cells
    note_top = table.rows[1].cells[3]
    note_bottom = table.rows[8].cells[3]
    note_top.merge(note_bottom)
    note_top.text = ""
    p = note_top.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, kahn_note, size=9, color=COLOR_GRAY_MED)

    # Total row
    total_cells = table.rows[9].cells
    total_cells[0].merge(total_cells[1])
    total_cells[0].text = ""
    p = total_cells[0].paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "Pre-vote total", bold=True, size=9.5, color=COLOR_BODY)
    set_cell_shading(total_cells[0], "F1F5F9")

    total_cells[2].text = ""
    p = total_cells[2].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "$404,171", bold=True, size=9.5, color=COLOR_BODY)
    set_cell_shading(total_cells[2], "F1F5F9")

    total_cells[3].text = ""
    set_cell_shading(total_cells[3], "F1F5F9")

    # Cell padding for all cells
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                f'  <w:top w:w="40" w:type="dxa"/>'
                f'  <w:left w:w="80" w:type="dxa"/>'
                f'  <w:bottom w:w="40" w:type="dxa"/>'
                f'  <w:right w:w="80" w:type="dxa"/>'
                f'</w:tcMar>'
            )
            tcPr.append(tcMar)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.2)
        row.cells[1].width = Inches(1.0)
        row.cells[2].width = Inches(1.0)
        row.cells[3].width = Inches(3.6)

    add_source(doc, "Source: JAHHS Project Expenditures ledger, BOC August 7, 2023 Agenda Packet, lines 2519-2526")

    # April 19, 2021
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(8)
    add_run(p, "April 19, 2021", bold=True, size=11, color=COLOR_BODY)
    add_run(p, ". Board approves $50M school project with USDA financing, 7-0.", size=11, color=COLOR_BODY)

    add_blockquote(doc,
        '"The current revaluation will also allow the county to realize an increase in property values but this will also result in an increase for tax payers. He stated that the County may have to adjust the rate to meet its obligations."',
        "Chair Kirby, BOC April 19, 2021 Minutes")

    add_body_paragraph(doc, "When Commissioner Kehayes raised concern about the debt-to-assessed-value ratio:")
    add_blockquote(doc,
        '"He thinks this should be addressed at the revaluation."',
        "County Manager Howard, BOC April 19, 2021 Minutes")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "August 2021", bold=True, size=11, color=COLOR_BODY)
    add_run(p, ". Truist offered a competitive financing rate of 2.26%.", size=11, color=COLOR_BODY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "September 20, 2021", bold=True, size=11, color=COLOR_BODY)
    add_run(p, ". BOC votes on Kahn contract. First formal Board vote. Approved 7-0.", size=11, color=COLOR_BODY)

    add_blockquote(doc,
        '"The school system did not have funds for the project when he was hired."',
        "Mr. Ott, BOC September 20, 2021 Minutes, lines 171-172")

    # ================================================================
    # 2022 (page break in HTML)
    # ================================================================
    add_page_break(doc)
    add_heading2(doc, "2022: Project Grows, Reappraisal Cycle Shortened")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "January 1, 2022", bold=True, size=11, color=COLOR_BODY)
    add_run(p, ". The 2022 reappraisal takes effect. Tax base increases 22.2% ($1.49B to $1.82B).", size=11, color=COLOR_BODY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "June 9, 2022", bold=True, size=11, color=COLOR_BODY)
    add_run(p, ". Special meeting on school construction. The Board approves a $75M project budget (Cummings sole dissenter). Key exchange on Kahn fees:", size=11, color=COLOR_BODY)

    add_blockquote(doc,
        '"Mr. Ott stated that at the time negotiating the construction manager at risk contract that the fee would not change. He stated now the fee has changed. He stated that he is concerned with the trust level that the fee would not change."',
        "Chair Kirby, BOC June 9, 2022 Special 9am Meeting Minutes")
    add_blockquote(doc,
        '"Mr. Ott was referring to the fee percentage not absolute dollar cash cost."',
        "Mr. Cram (M.B. Kahn), same meeting")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "June 9, 2022", bold=True, size=11, color=COLOR_BODY)
    add_run(p, ". Board sets tax rate. Revenue-neutral rate: $0.635. County Manager Howard proposed $0.73. Board adopted $0.665, using $1.15M from fund balance. Howard advises:", size=11, color=COLOR_BODY)

    add_blockquote(doc,
        '"If they spend fund balance in the next fiscal year, these monies will have to be located in the budget next fiscal year."',
        "County Manager Howard, BOC June 9, 2022 Minutes")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "August 15, 2022", bold=True, size=11, color=COLOR_BODY)
    add_run(p, ". Board votes 7-0 to shorten reappraisal cycle from 8 years to 4 years. No public hearing held on the cycle change.", size=11, color=COLOR_BODY)

    add_blockquote(doc,
        '"By doing this would benefit taxpayers as the valuation would better follow the market and reduce sticker shock."',
        "Tax Administrator Melissa Radke, BOC August 15, 2022 Minutes")

    add_source(doc, "This cycle change was adopted 5 months after the 2022 reappraisal and 14 months after the USDA loan. The Board's discussion contains no reference to school debt, the USDA loan, or revenue needs.")

    # 2023
    add_heading2(doc, "2023: GMP Approved, PNC Financing")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "February 6, 2023", bold=True, size=11, color=COLOR_BODY)
    add_run(p, ". Vincent Valuations contracted for 2026 reappraisal ($364,687). McLaughlin recused. Lawrence voted yes despite owning 47 parcels subject to the reappraisal.", size=11, color=COLOR_BODY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "September 5, 2023", bold=True, size=11, color=COLOR_BODY)
    add_run(p, ". Board approves $82.4M Guaranteed Maximum Price (GMP). Chair Kirby discloses his brother's company received a subcontract on the project. Kirby voted yes.", size=11, color=COLOR_BODY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "November 20, 2023", bold=True, size=11, color=COLOR_BODY)
    add_run(p, ". PNC Bank financing: $34,998,500 at a 9% rate floor. No competitive solicitation documented. In August 2021, Truist had offered a competitive rate of 2.26%. The PNC rate floor of 9% represents a significantly higher cost of financing. No competitive solicitation for the PNC financing is documented.", size=11, color=COLOR_BODY)

    # 2024-2025
    add_heading2(doc, "2024-2025: Rate Climbs, Fund Balance Drawn")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "June 2024", bold=True, size=11, color=COLOR_BODY)
    add_run(p, ". Tax rate raised to $0.695. School project now at approximately $85M.", size=11, color=COLOR_BODY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "June 2025", bold=True, size=11, color=COLOR_BODY)
    add_run(p, ". Board uses $670K in Timbermill wind farm revenue and additional fund balance to avoid another rate increase. Howard's 2022 advisement about fund balance spending has materialized.", size=11, color=COLOR_BODY)

    # 2026
    add_heading2(doc, "2026: The Reappraisal Arrives")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "September 15, 2025", bold=True, size=11, color=COLOR_BODY)
    add_run(p, ". 2026 reappraisal presented. Increases range from 15% to 161%.", size=11, color=COLOR_BODY)

    add_blockquote(doc,
        '"The market will cause those who live on a fixed income and are not buying or selling to be priced out of their homes."',
        "Commissioner Cummings, BOC September 15, 2025 Minutes")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "October 20, 2025", bold=True, size=11, color=COLOR_BODY)
    add_run(p, ". Schedule of Values adopted. Zero public comments. Passed 6-0.", size=11, color=COLOR_BODY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "January 1, 2026", bold=True, size=11, color=COLOR_BODY)
    add_run(p, ". Reappraisal takes effect. County-wide increase estimated at 38-42%.", size=11, color=COLOR_BODY)

    # ================================================================
    # BY THE NUMBERS (page break)
    # ================================================================
    add_page_break(doc)
    add_heading1(doc, "By the Numbers")

    create_table(doc,
        ["What Changed", "Before (2021)", "After"],
        [
            ["School project cost", "$50 million (approved Apr 2021)", "~$85 million (reported cost)"],
            ["Voter approval?", "Yes (bond, rejected 2018)", "No (USDA installment)"],
            ["County debt (pre-Holmes)", "$7M (retiring by 2025)", "Retired"],
            ["Holmes HS debt (new)", "$50M (approved)", "~$85M (reported)"],
            ["Holmes HS annual debt service (estimated)", "~$2M/yr (Davenport projection at $50M)", "~$3.4 million/year (estimated)"],
            ["School debt as % of property tax levy (estimated)", "~17.9%", "~24.3%"],
            ["Tax rate", "$0.755 (FY2021-22)", "$0.695 (FY2025-26)"],
            ["Tax base", "$1.49 billion", "$2.00 billion"],
            ["Property tax revenue", "~$11.2M", "~$14.0M (+25%)"],
            ["Fund balance", "27.53% (2022)", "29.50% (Finance Officer, Feb 2026) / 58.13% (audit)"],
            ["Reappraisal cycle", "8 years", "4 years"],
        ],
        col_widths=[Inches(2.8), Inches(2.1), Inches(2.1)],
    )

    # Table note
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    add_run(p, 'How "School debt as % of property tax levy" is calculated: ', bold=True, size=9, color=COLOR_GRAY_MED, font_name="Calibri")
    add_run(p, 'The estimated annual debt service payment is divided by the total property tax revenue for that period. At the $50M project level, Davenport projected approximately $2M per year in debt service; the property tax levy at that time was approximately $11.2M, producing $2M / $11.2M = 17.9%. At the reported $85M project level, estimated debt service is approximately $3.4M per year; the current property tax levy is approximately $14M, producing $3.4M / $14M = 24.3%. Both figures represent the school debt service portion only, not total school funding (which includes operational expenses). These are estimates based on publicly available data; the actual debt service schedule has not been made publicly available.', size=9, color=COLOR_GRAY_MED, font_name="Calibri")

    # ================================================================
    # KAHN FEE TABLE
    # ================================================================
    add_heading1(doc, "M.B. Kahn Fee & Project Cost Changes")

    create_table(doc,
        ["Metric", "Original Scope", "As Costs Grew"],
        [
            ["Construction cost base", "$40,000,000", "$64,744,592"],
            ["CMAR fee (3.5%)", "$1,400,000", "$2,266,061"],
            ["LS3P architect fee (5.75%)", "$2,300,000", "$3,722,814"],
            ["Total project budget", "$50,000,000", "$75,000,000 to ~$85,000,000"],
            ["Total Kahn payments (through Aug 2023)", "$1,407,511", ""],
        ],
        col_widths=[Inches(2.8), Inches(2.1), Inches(2.1)],
    )

    add_source(doc, "Sources: BOC June 9, 2022 Special Meeting Minutes; JAHHS Project Expenditures, BOC Aug 7, 2023 Agenda Packet")

    # ================================================================
    # GUIDANCE TABLE
    # ================================================================
    add_heading1(doc, "Where Did the Guidance Come From?")

    create_table(doc,
        ["Decision", "Guidance Source", "What They Said", "Their Interest"],
        [
            ["No referendum needed", "Rick Ott, M.B. Kahn (contractor)", '"No that is not required"', "Kahn was hired to build the school"],
            ['"No impact on tax rate"', "Davenport & Associates (advisor)", "Based on $50M project", "Paid up to $30K for the engagement"],
            ['"Reduce sticker shock"', "Tax Administrator Melissa Radke", '"Benefit taxpayers"', "Presented after USDA loan approved"],
            ["Values are accurate", "Vincent Valuations (contractor)", "Median 99.04%", "Paid $378K (2022) and $365K (2026)"],
        ],
        col_widths=[Inches(1.3), Inches(1.7), Inches(2.1), Inches(1.7)],
    )

    # ================================================================
    # TAX IMPACT (page break)
    # ================================================================
    add_page_break(doc)
    add_heading1(doc, "What the Reappraisal Means for Your Tax Rate")

    add_heading2(doc, "What Has Already Happened to Your Tax Bill")
    add_body_paragraph(doc, "Example: a home assessed at $200,000 before the 2022 reappraisal.")

    create_table(doc,
        ["Year", "Assessed Value", "Tax Rate", "Annual Bill", "Change from FY2021-22"],
        [
            ["FY2021-22 (before reappraisal)", "$200,000", "$0.755", "$1,510", "n/a"],
            ["FY2022-23 (after 2022 reappraisal)", "$244,400", "$0.665", "$1,625", "+$115 (+7.6%)"],
            ["FY2024-25 (rate increase)", "$244,400", "$0.695", "$1,699", "+$189 (+12.5%)"],
        ],
        col_widths=[Inches(2.2), Inches(1.2), Inches(1.0), Inches(1.1), Inches(1.3)],
    )

    add_body_paragraph(doc, "Your bill has already increased 12.5% since FY2021-22, even though the rate dropped from $0.755 to $0.695.")

    add_heading2(doc, "2026 Scenarios Based on Board's Documented Pattern")
    add_body_paragraph(doc, "The same home after the 2026 reappraisal (est. 40% increase, new value ~$342,000):")

    create_table(doc,
        ["Scenario", "Rate", "Annual Bill", "Change from Today", "Cumulative since FY2021-22"],
        [
            ["Revenue-neutral (RNR)", "~$0.498", "~$1,704", "+$5 (+0%)", "+$194 (+13%)"],
            ["Board pattern: 5% above RNR", "~$0.523", "~$1,789", "+$91 (+5%)", "+$279 (+18%)"],
            ["Board pattern: 10% above RNR [LIKELY]", "~$0.548", "~$1,874", "+$176 (+10%)", "+$364 (+24%)"],
            ["Manager pattern: 15% above RNR", "~$0.573", "~$1,960", "+$261 (+15%)", "+$450 (+30%)"],
        ],
        col_widths=[Inches(2.2), Inches(1.0), Inches(1.1), Inches(1.2), Inches(1.3)],
        highlight_rows={2},
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    add_run(p, "Scenarios based on documented Board behavior. In 2022, the Board adopted 4.7% above RNR ($0.665 vs $0.635 RNR). The County Manager proposed 15% above. The Board has never adopted a revenue-neutral rate after a reappraisal.", size=9, color=COLOR_GRAY_MED, font_name="Calibri")

    # ================================================================
    # NC COMPARISON (page break)
    # ================================================================
    add_page_break(doc)
    add_heading2(doc, "How Does This Compare to Other NC Counties?")

    add_body_paragraph(doc, "A 2024 UNC School of Government study of 200 reappraisals across all 100 NC counties (2008-2023) found:")

    add_bullet(doc, " of NC counties adopted rates above revenue-neutral after reappraisals.", bold_prefix="67%")
    add_bullet(doc, " adopted above revenue-neutral, averaging 6% above.", bold_prefix="When tax bases grew, 92%")
    add_bullet(doc, "Chowan County adopted 4.7% above revenue-neutral in 2022.")
    add_bullet(doc, "County Manager Howard confirmed in April 2019 that the Board chose not to go revenue-neutral after the 2014 reappraisal.")

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "Chowan County has never adopted a revenue-neutral rate after a reappraisal", bold=True, size=11, color=COLOR_BODY)
    add_run(p, ", based on the documented record.", size=11, color=COLOR_BODY)

    add_source(doc, "Sources: UNC School of Government, McLaughlin (2024); BOC April 17, 2019 Minutes.")

    # ================================================================
    # WHAT YOU CAN DO (page break)
    # ================================================================
    add_page_break(doc)
    add_heading1(doc, "What You Can Do")

    add_heading2(doc, "Before May 11, 2026: If Your Assessment Contains Errors")
    add_bullet(doc, "Review your assessment notice for data errors (square footage, year built, bedrooms, lot size)")
    add_bullet(doc, "File an appeal with the Board of Equalization and Review if evidence supports a lower value")
    add_bullet(doc, "Bring: comparable sales data, a private appraisal, photos, or error documentation")
    add_bullet(doc, "You do not need an attorney")

    add_heading2(doc, "Before the Budget Hearing (May/June 2026): This Is Where the Rate Is Set")
    add_bullet(doc, "Ask the county to publish the revenue-neutral tax rate before the hearing")
    add_bullet(doc, "Attend the budget hearing and speak during public comment")
    add_bullet(doc, "Ask: How much of the tax rate goes to school debt service?")
    add_bullet(doc, "Ask: What would the rate be without the school debt?")
    add_bullet(doc, "Ask: Was the Davenport financial analysis updated when the project cost increased from $50M to $85M?")

    add_note(doc, "Note", [
        "Zero public comments were received at both the 2022 and 2026 Schedule of Values hearings. The budget hearing should not be zero again.",
    ])

    add_heading2(doc, "Tax Relief Programs")
    add_body_paragraph(doc, "If you are 65+, disabled, or on a fixed income, NC law provides programs that may reduce your tax bill:")

    add_bullet(doc, " (G.S. \u00a7 105-277.1)", bold_prefix="Elderly/Disabled Exclusion")
    add_bullet(doc, " (G.S. \u00a7 105-277.1B)", bold_prefix="Circuit Breaker")
    add_bullet(doc, " (G.S. \u00a7 105-277.1C)", bold_prefix="Tax Deferral")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_run(p, "Visit ", size=11, color=COLOR_BODY)
    add_run(p, "chowancounty-nc.gov", bold=True, size=11, color=COLOR_BODY)
    add_run(p, " for contact information.", size=11, color=COLOR_BODY)

    # ================================================================
    # SCOPE OF REVIEW
    # ================================================================
    add_heading1(doc, "Scope of Review")

    add_heading2(doc, "What Was Reviewed")
    add_bullet(doc, " from the Chowan County Board of Commissioners (meeting minutes, agenda packets, retreat minutes, special meeting minutes, CIP Committee minutes, Joint Committee minutes)", bold_prefix="480 documents")
    add_bullet(doc, " 2019 through March 2026", bold_prefix="Coverage:")
    add_bullet(doc, " All documents downloaded from chowancounty-nc.gov", bold_prefix="Source:")
    add_bullet(doc, " Every document SHA-256 hashed at acquisition. Full hash log available upon request.", bold_prefix="Verification:")
    add_bullet(doc, " Verified against county website April 2, 2026. Collection is complete relative to what the county has made publicly available.", bold_prefix="Completeness:")

    add_heading2(doc, "Records Needed to Complete This Analysis")

    foia_items = [
        ("Board of Education meeting minutes. The BOE is a separate body. BOC minutes reference that Kahn was presented as a \"request from Edenton-Chowan Board of Education.\" The BOE minutes may contain the Kahn hiring authorization.",
         "Needed from: Edenton-Chowan Board of Education"),
        ("M.B. Kahn CMAR contract. The executed agreement including fee schedule, scope, amendments, and change orders.",
         "Needed from: Chowan County Manager's Office"),
        ("Kahn qualifications-based selection records. Under G.S. 143-128.1, CMAR uses qualifications-based selection. No RFQ, selection committee records, or evaluation summaries were found in any BOC document.",
         "Needed from: Chowan County Manager's Office and/or Board of Education"),
        ("USDA loan agreement and amortization schedule. The executed loan terms. Did USDA approve the cost increases from $50M to $85M?",
         "Needed from: Chowan County Manager's Office / USDA Rural Development"),
        ("Davenport & Associates updated financial projections. The original projection was \"no impact on current tax rate\" at $50M. Was this updated when costs reached $85M?",
         "Needed from: Chowan County Manager's Office"),
        ("PNC Bank financing selection records. PNC financing was $34.9M at a 9% floor. Truist offered 2.26% in 2021. No competitive quotes documented.",
         "Needed from: Chowan County Manager's Office"),
    ]

    for desc, target in foia_items:
        add_foia_item(doc, desc, target)

    add_heading2(doc, "Note on County Audits")
    add_body_paragraph(doc, "The county's annual financial audit verifies that dollar amounts are correctly recorded. It does not audit processes or programs. The procedural questions raised in this report are outside the scope of a financial statement audit.")

    # ================================================================
    # APPENDIX (page break)
    # ================================================================
    add_page_break(doc)
    add_heading1(doc, "Appendix: Key Figures Referenced in This Report")

    create_table(doc,
        ["Name", "Role", "Relevance to This Report"],
        [
            ["Patti Kersey", "BOC Chair (2019-2020)", "Presided over referendum cancellation and early USDA discussions"],
            ["Bob Kirby", "Commissioner, then Chair (2021-present)", "Asked if referendum needed; linked revaluation to debt capacity; confronted Kahn on fee increase"],
            ["Kevin Howard", "County Manager", 'Proposed budgets; advised about fund balance; stated debt ratio would be "addressed at the revaluation"'],
            ["Rick Ott", "Construction Manager, M.B. Kahn", "Advised Board no referendum needed for USDA. Present from Jan 2020, 20 months before contract vote."],
            ["Mr. Cram", "M.B. Kahn representative", "Distinguished fee percentage from dollar amount in response to Kirby's confrontation"],
            ["Ron Cummings", "Commissioner", "Sole consistent dissenter on cost increases; raised concern about fixed-income residents"],
            ["Larry McLaughlin", "Commissioner", 'Objected 2019 bond process "seems rushed"; recused on Vincent contract'],
            ["Commissioner Kehayes", "Commissioner", "Raised concern about debt-to-assessed-value ratio"],
            ["Ryan Vincent", "Owner, Vincent Valuations", "Conducted 2022 and 2026 reappraisals"],
            ["Dr. Sasscer", "Superintendent", "Confirmed referendum timeline; proposed capping Kahn fee at $50M base"],
            ["Chris Hill", "Tax Administrator", "Presented reappraisal data"],
            ["Cathy Smith", "Finance Officer", "Presented budget and financial data"],
            ["Melissa Radke", "Former Tax Administrator", "Presented resolution to shorten reappraisal cycle"],
            ["Davenport & Associates", "County's financial advisor", 'Projected "no impact on current tax rate" for $50M project'],
        ],
        col_widths=[Inches(1.5), Inches(2.0), Inches(3.3)],
    )

    # ================================================================
    # TERMS OF USE
    # ================================================================
    add_heading1(doc, "Terms of Use")

    tos_items = [
        "This analysis is produced from public records and is provided for informational purposes only.",
        "It does not constitute legal, financial, or tax advice.",
        'Information is provided "as is" without warranty of any kind.',
        "If any information is found to be inaccurate, it will be updated to match official records upon notification.",
        "This is a product of the written public record, simplified for citizen accessibility.",
    ]
    for item in tos_items:
        add_bullet(doc, item, size=9, indent=Inches(0.4))

    # ================================================================
    # FOOTER
    # ================================================================
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    set_paragraph_border_bottom(p, color="E2E8F0", sz="4")
    # Empty paragraph for the border

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    add_run(p, "GovInsights", bold=True, size=9, color=COLOR_GRAY_MED)
    add_run(p, " is a project of Ropewalk Technologies LLC.", size=9, color=COLOR_SOURCE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    add_run(p, "All data sourced from official Chowan County Board of Commissioner meeting minutes and agenda packets, with supporting data from the U.S. Census Bureau, NC Department of Public Instruction, NC Department of Revenue, UNC School of Government, and NC General Statutes.", size=9, color=COLOR_SOURCE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    add_run(p, "This document was prepared with the assistance of AI for data compilation and formatting. All factual claims are sourced to government records.", size=9, color=COLOR_SOURCE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    add_run(p, "Contact: govinsights.io/contact", size=9, color=COLOR_SOURCE)

    # Save
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
